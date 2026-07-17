"""
report_generator.py — renders a calculated project + its ledger audit trail
into a registry-format (Verra / Puro.earth / Isometric-style) Monitoring
Report .docx.

This is a RENDERING layer only — it doesn't recompute anything. Every
number here was already computed by vm0044_engine.py and already committed
to the ledger by views/projects.py; this just lays it out the way a VVB
expects to review it: project description, fixed (ex-ante) parameters,
monitored (ex-post) per-batch data, the quantification results, QA/QC,
and a live audit-trail excerpt with the ledger's public key attached.

If you need to match one specific registry's exact template headings,
edit the section titles/order below (or fork this file) — the underlying
data it pulls from (project_store + the ledger) doesn't change.
"""

import io

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DEEP_BLUE = RGBColor(0x12, 0x3F, 0x5E)
ACCENT_GREEN = RGBColor(0x2E, 0x9E, 0x5B)


# ---------------------------------------------------------------------------
# small docx helpers
# ---------------------------------------------------------------------------

def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DEEP_BLUE
    return h


def _style_table(table):
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def _df_to_table(doc, df: pd.DataFrame, max_rows: int = 200):
    """Adds a DataFrame as a Word table. Truncates very long tables with a note
    rather than silently dropping rows or blowing up the document."""
    if df is None or len(df) == 0:
        doc.add_paragraph("None.")
        return
    truncated = len(df) > max_rows
    view = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(view.columns))
    _style_table(table)
    for i, col in enumerate(view.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(view.columns):
            val = row[col]
            cells[i].text = "—" if pd.isna(val) else str(val)
    if truncated:
        p = doc.add_paragraph(
            f"(showing first {max_rows} of {len(df)} rows — export the full CSV "
            f"from the Projects tab for the complete table)"
        )
        p.runs[0].italic = True


def _kv_table(doc, pairs):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for k, v in pairs:
        row = table.add_row().cells
        row[0].text = str(k)
        if row[0].paragraphs[0].runs:
            row[0].paragraphs[0].runs[0].bold = True
        row[1].text = "—" if v in (None, "") else str(v)


# ---------------------------------------------------------------------------
# main entry point
# ---------------------------------------------------------------------------

def build_docx_report(
    project_id: str,
    project: dict,
    config_effective: dict,
    report_meta: dict,
    chain_ok: bool,
    chain_issues: list,
    trail_records: list,
    public_key_pem: bytes,
) -> bytes:
    """
    project           — the dict stored by project_store.save_project()
                         (meta / lat / lon / result / config_overrides)
    config_effective  — vm0044_engine.load_config(project.get("config_overrides")),
                         i.e. the methodology defaults actually used, merged
                         with any project-specific override
    report_meta       — {registry, methodology, period_start, period_end,
                         report_version, prepared_by, report_date, notes}
    chain_ok/chain_issues — result of ledger.verify_chain(), run fresh at
                         report-generation time (not cached/asserted)
    trail_records     — this project's own ledger entries (list of dicts),
                         already filtered by the caller
    public_key_pem    — ledger.export_public_key_pem(), so a VVB has
                         everything needed to verify signatures independently

    Returns the .docx file as raw bytes, ready for a Streamlit download_button.
    """
    result = project["result"]
    ledger_df = result["ledger"]
    summary = result["summary"]
    meta = project["meta"]

    doc = Document()

    # ---------------- Cover page ----------------
    title = doc.add_heading("GHG Monitoring Report", level=0)
    for run in title.runs:
        run.font.color.rgb = DEEP_BLUE
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{meta.get('Project Name', project_id)}  ·  {project_id}")
    run.font.size = Pt(15)
    run.font.color.rgb = ACCENT_GREEN

    doc.add_paragraph()
    _kv_table(doc, [
        ("Registry", report_meta["registry"]),
        ("Methodology", report_meta["methodology"]),
        ("Monitoring period", f"{report_meta['period_start']} to {report_meta['period_end']}"),
        ("Report version", report_meta["report_version"]),
        ("Prepared by", report_meta["prepared_by"]),
        ("Report date", report_meta["report_date"]),
    ])
    if report_meta.get("notes"):
        doc.add_paragraph()
        note_p = doc.add_paragraph("Notes for the VVB: " + report_meta["notes"])
        note_p.runs[0].italic = True

    doc.add_page_break()

    # ---------------- 1. Project description ----------------
    _add_heading(doc, "1. Project Description", level=1)
    _kv_table(doc, [
        ("Project ID", project_id),
        ("Project Name", meta.get("Project Name", "—")),
        ("Facility Address", meta.get("Facility Address", "—")),
        ("Facility Coordinates", f"{project['lat']:.4f}, {project['lon']:.4f}"),
        ("Pyrolysis Reactor(s)", meta.get("Pyrolysis Reactor(s)", "—")),
        ("Project Start Date", meta.get("Project Start Date", "—")),
        ("Project Supervisor", meta.get("Project Supervisor", "—")),
        ("Registry", meta.get("Registry", "—")),
    ])

    # ---------------- 2. Quantification approach ----------------
    _add_heading(doc, "2. GHG Quantification Approach", level=1)
    doc.add_paragraph(
        "Net GHG removals are quantified per batch as: carbon stock in biochar "
        "(dry, ash-corrected mass x fixed carbon fraction FCp x permanence factor "
        "PRde,k x 44/12), less project emissions (pyrolysis energy use + fugitive "
        "methane, low-technology pathway per VM0044 Equation 9) and less transport "
        "leakage beyond the applicable one-way distance threshold. Batches failing "
        "the H:Corg soil-application eligibility threshold (Applicability Condition "
        "10b) are excluded from the credited pool entirely — net_tCO2e is set to "
        "zero and the batch is flagged INELIGIBLE — rather than assumed eligible by "
        "default."
    )

    # ---------------- 3. Ex-ante parameters ----------------
    _add_heading(doc, "3. Data and Parameters Fixed at Validation (Ex-Ante)", level=1)
    doc.add_paragraph(
        "Methodology default values used for this project's calculation run. Any "
        "project-specific override applied at run time is noted below the tables."
    )
    ante_df = pd.DataFrame([
        ("CO2 per unit C (44/12)", config_effective["constants"]["co2_per_c"], "—", "Stoichiometric constant"),
        ("GWP of CH4 (AR5, 100yr)", config_effective["constants"]["gwp_ch4"], "—", "IPCC AR5, VM0044 Sec 9.1"),
        ("Default Fe, low-tech kiln", config_effective["constants"]["fe_default_low_tech"], "tCH4/t biochar", "Cornelissen et al. 2016"),
        ("H:Corg eligibility threshold", config_effective["constants"]["h_corg_eligibility_threshold"], "—", "Applicability Condition 10b"),
        ("Transport leakage distance threshold", config_effective["constants"]["transport_leakage_threshold_km"], "km one-way", "CDM Tool 16"),
        ("Diesel emission factor", config_effective["emission_factors"]["diesel_ef_kgco2_per_l"], "kgCO2/L", "IPCC standard"),
        ("Grid emission factor", config_effective["emission_factors"]["grid_ef_kgco2_per_kwh"], "kgCO2/kWh", "Project/country grid EF"),
        ("Transport emission factor", config_effective["emission_factors"]["transport_ef_kgco2_per_tkm"], "kgCO2/t-km", "CDM Tool 12 / local factor"),
        ("Permanence factor, low-tech default (PRde,k)", config_effective["permanence_factor"]["low_tech_default"], "—", "VM0044 Table 4"),
    ], columns=["Parameter", "Value", "Unit", "Source / basis"])
    _df_to_table(doc, ante_df)

    fc_rows = []
    for feed, procs in config_effective["fixed_carbon_defaults"].items():
        for proc, val in procs.items():
            fc_rows.append({"Feedstock category": feed, "Process": proc, "Default FCp": val})
    doc.add_paragraph()
    fc_label = doc.add_paragraph("Default fixed carbon fraction (FCp) by feedstock and process:")
    fc_label.runs[0].bold = True
    _df_to_table(doc, pd.DataFrame(fc_rows))

    if project.get("config_overrides"):
        p = doc.add_paragraph()
        p.add_run("Project-specific overrides applied to the defaults above: ").bold = True
        p.add_run(str(project["config_overrides"]))

    # ---------------- 4. Ex-post monitored data ----------------
    _add_heading(doc, "4. Data and Parameters Monitored (Ex-Post, Per Batch)", level=1)
    doc.add_paragraph(
        "Values measured or recorded per batch during the monitoring period, as "
        "submitted in the project's data tables (Suppliers & Feedstocks, Pyrolysis "
        "Batches, Lab_Data)."
    )
    monitored_cols = [c for c in [
        "batch_id", "feedstock_category", "process", "tech_level",
        "biochar_output_wet_kg", "moisture_pct", "ash_pct",
        "lab_fc", "fc_used", "h_corg", "tprod_c",
        "diesel_l", "electricity_kwh", "transport_km_one_way",
    ] if c in ledger_df.columns]
    _df_to_table(doc, ledger_df[monitored_cols])

    # ---------------- 5. Quantification results ----------------
    _add_heading(doc, "5. GHG Emission Reduction / Removal Quantification", level=1)
    calc_cols = [c for c in [
        "batch_id", "dry_biochar_mass_t", "soil_eligible", "carbon_stock_t",
        "co2e_removed_gross", "pe_energy_tCO2e", "pe_methane_tCO2e",
        "leakage_transport_tCO2e", "net_tCO2e", "flag",
    ] if c in ledger_df.columns]
    _df_to_table(doc, ledger_df[calc_cols])

    doc.add_paragraph()
    summary_label = doc.add_paragraph("Project-level summary:")
    summary_label.runs[0].bold = True
    _kv_table(doc, [
        ("Total batches", summary["total_batches"]),
        ("Eligible batches", summary["eligible_batches"]),
        ("Ineligible batches", summary["ineligible_batches"]),
        ("Gross CO2e removed (tCO2e)", summary["gross_co2e_removed"]),
        ("Total project emissions (tCO2e)", summary["total_project_emissions"]),
        ("Total leakage (tCO2e)", summary["total_leakage"]),
        ("NET tCO2e (this calculation run)", summary["net_tCO2e_project"]),
    ])

    # ---------------- 6. QA/QC ----------------
    _add_heading(doc, "6. QA/QC Procedures and Results", level=1)
    doc.add_paragraph(
        "Automated checks applied to every batch: (a) soil-application eligibility "
        "gate on H:Corg, (b) mass-balance check of applied vs. produced dry biochar "
        "mass (1% tolerance), (c) a missing-lab-data flag where H:Corg was not "
        "measured, and — where satellite sourcing data is available — (d) an "
        "over-harvest and classification-confidence cross-check at the feedstock "
        "delivery level against the GEE biomass estimate for that plot."
    )
    _kv_table(doc, [
        ("Batches flagged — mass balance", summary["batches_flagged_mass_balance"]),
        ("Batches flagged — missing lab data", summary["batches_flagged_missing_lab"]),
    ])
    if len(result["qa_flags"]) > 0:
        doc.add_paragraph()
        flag_label = doc.add_paragraph("Flagged batches:")
        flag_label.runs[0].bold = True
        _df_to_table(doc, result["qa_flags"])
    else:
        doc.add_paragraph("No batch-level QA issues found in this calculation run.")

    if "satellite_sourcing_qa" in result:
        doc.add_paragraph()
        sat_label = doc.add_paragraph("Satellite sourcing cross-check (feedstock delivery vs. GEE biomass estimate):")
        sat_label.runs[0].bold = True
        _df_to_table(doc, result["satellite_sourcing_qa"])

    # ---------------- 7. Audit trail ----------------
    _add_heading(doc, "7. Data Management and Audit Trail", level=1)
    doc.add_paragraph(
        "Every batch calculation and project-level run above is committed to an "
        "append-only, hash-chained, Ed25519-signed ledger at the moment it is "
        "calculated — not reconstructed after the fact. Any retroactive edit to a "
        "past entry breaks the hash chain from that point forward and invalidates "
        "its signature. Both checks are independently reproducible by a third party "
        "who has only the exported ledger file (.jsonl) and the public key below — "
        "no access to this application, its server, or its database is required."
    )
    _kv_table(doc, [
        ("Chain integrity check (run at report generation time)",
         "PASSED — no tampering detected" if chain_ok else f"FAILED — {len(chain_issues)} issue(s) found"),
        ("Ledger entries for this project", len(trail_records)),
    ])
    if trail_records:
        seqs = [r["seq"] for r in trail_records]
        doc.add_paragraph(f"Sequence numbers {min(seqs)}\u2013{max(seqs)} of the full ledger chain.")
        trail_df = pd.DataFrame(trail_records)[["seq", "record_type", "record_id", "timestamp", "entry_hash"]]
        _df_to_table(doc, trail_df)
    if not chain_ok:
        doc.add_paragraph()
        warn_p = doc.add_paragraph()
        warn_p.add_run("Chain issues detected — resolve before submission:").bold = True
        for issue in chain_issues:
            doc.add_paragraph(str(issue), style="List Bullet")

    doc.add_paragraph()
    key_label = doc.add_paragraph("Ledger signing public key (Ed25519, PEM) — for independent signature verification:")
    key_label.runs[0].bold = True
    key_para = doc.add_paragraph(public_key_pem.decode("utf-8"))
    for run in key_para.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(8)

    # ---------------- 8. Declaration ----------------
    _add_heading(doc, "8. Declaration", level=1)
    doc.add_paragraph(
        "I confirm that the data and calculations presented in this report reflect "
        "this project's monitored performance for the stated monitoring period, "
        "calculated per the methodology and parameters stated above, and that the "
        "audit trail referenced in Section 7 has not been altered since each entry "
        "was committed."
    )
    doc.add_paragraph()
    _kv_table(doc, [
        ("Prepared by", report_meta["prepared_by"]),
        ("Date", report_meta["report_date"]),
        ("Signature", "_______________________________"),
    ])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
